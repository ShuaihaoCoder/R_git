from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r"C:\Users\PC\Desktop\R_git\YieldCurve_Shiny")
TARGET_IMAGE = PROJECT / "sampleUI" / "04_carry_roll.png"
CURRENT_IMAGE = PROJECT / "current_carry_roll_applied.png"
DOCX_PATH = PROJECT / "Carry_Roll_UI_Comparison_Plan.docx"


INK = "10263B"
BLUE = "1F5F8B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
LINE = "D7E1EB"
GREEN = "0E8A7A"
RED = "D94B45"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, bold=False, color=INK, size=9):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    style_run(run, bold=True, color=BLUE if level <= 2 else INK, size={1: 16, 2: 13, 3: 11}.get(level, 11))
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_body(doc, text, bold=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    style_run(run, bold=bold, size=9.5)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    style_run(run, size=9)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    style_run(run, color="66788C", size=8)
    return paragraph


def add_callout(doc, title, lines, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [9.0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    style_run(run, bold=True, color=INK, size=9.5)
    for line in lines:
      p = cell.add_paragraph()
      p.paragraph_format.space_after = Pt(2)
      r = p.add_run(line)
      style_run(r, size=8.7)
    doc.add_paragraph()


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [1.75, 7.25])
    for idx, title in enumerate(["字段", "内容"]):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(title)
        style_run(r, bold=True, size=8.5)
    for key, value in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        r0 = cells[0].paragraphs[0].add_run(key)
        style_run(r0, bold=True, size=8.3)
        r1 = cells[1].paragraphs[0].add_run(value)
        style_run(r1, size=8.2)
    doc.add_paragraph()


def add_gap_table(doc):
    headers = ["编号", "模块", "目标 UI 表达", "当前 UI 截图状态", "当前代码状态", "差距判断", "推荐修改", "实现风险", "用户补充/决策"]
    rows = [
        ["1", "顶部导航", "深 navy 顶栏，Carry & Roll active，右侧显示 Market Data Live、通知、用户头像。", "当前为深 navy 顶栏，Carry & Roll active；右侧为 Local RDS Loaded / AB，非 live 数据。", "navbarPage + styles.css navbar 伪元素显示本地 RDS 状态。", "视觉接近，但文案/状态语义不同。", "保留本地真实状态，不伪造 Market Data Live；可把右侧状态做成目标图布局但显示 Local RDS。", "低。注意不要把本地数据伪装为实时 Bloomberg。", ""],
        ["2", "左侧 Sidebar", "按 Data Source、Curve、As of Date、Tenor Range、Hold Period、Direction、DV01、按钮、Status 排列。", "当前为 Carry Controls，分 Single trade source、Single trade、Curve trade structure 两大块；字段多且滚动较长。", "app.R 中 Carry tab 使用 source_controls('carry') 和 source_controls('trade')，再放 numeric/select inputs。", "信息架构差距大。目标把交易员首屏操作合并为一个更短控制栏。", "重排 Carry sidebar；Curve Trade 默认复用 Carry 数据源，保留隐藏/兼容 trade inputs。", "中。会影响 input id、测试和 applied-result pending 逻辑。", ""],
        ["3", "页面标题区", "Carry & Roll Analysis，下方 Source、Curve、Effective Date、Hold Basis metadata。", "当前标题是 Carry & Roll，副标题是说明句；metadata 在 carry_banner 中出现在 section 内。", "page_header() + uiOutput('carry_banner')。", "目标标题更像交易工作台，当前更像文档式模块页。", "新增 carry title/meta block，把 carry_banner 信息提升到主区顶部。", "低。缺失 Bloomberg metadata 时显示 Local RDS / NA。", ""],
        ["4", "KPI Cards", "6 张横向卡：Carry、Roll、Total、P&L、Trade Mode、Curve Quality，带正负色和 subtitle。", "当前 Single Trade section 内只有 4 张卡：Carry、Roll、Total / P&L、Direction。", "renderText: carry_value、roll_value、total_value、carry_direction_value。", "缺 P&L 独立、Trade Mode、Curve Quality；层级位置不对。", "拆分 total 和 P&L outputs，新增 trade mode / quality outputs；KPI 放到标题下。", "中。Curve Quality 只能用 RMSE/proxy 推导，不可伪造 Good。", ""],
        ["5", "Single Trade 图", "目标是多期限 Carry/Roll bar + Total line，标题为 Single Trade Decomposition (3M)。", "当前首屏右侧小图只显示单笔 Carry、Roll、Total 三根柱。多期限图在下一行 carry_stacked_plot，且用 facet 按 hold 拆开。", "carry_component_plot 用单笔三柱；carry_stacked_plot 用 matrix + facet_wrap。", "图表语义差距大。", "把 carry_stacked_plot 的 3M hold 版本提升为目标主图：teal carry bar、red roll bar、navy total line。", "中。需要保证 hold filter、legend、tooltip 和 x 轴顺序。", ""],
        ["6", "Spot Curve 图", "目标图右上，点线图，x 轴 1M 到 30Y，标题含曲线名。", "当前 Spot Curve 已在首屏右侧，可见但坐标/留白/标题还是 ggplotly 默认风格。", "carry_spot_plot 使用 ggplot + ggplotly。", "功能存在，视觉需要收敛。", "改为显式 Plotly trace 或调 ggplotly layout：legend、margin、modebar、sparse tenor tick。", "低到中。Plotly modebar 容易覆盖标题。", ""],
        ["7", "Heatmap", "目标左下是矩阵热力图，Rows=Hold Period，Columns=Tenor，带底部红-白-绿 colorbar。", "当前 Heatmap 在右侧小卡，左侧反而是很高的 stacked/facet 图。", "carry_heatmap 已基于 matrix 输出；carry_matrix 另有 DT 表。", "数据和图表已有，但位置和尺寸不对。", "把 carry_heatmap 放到左下大卡，隐藏或降级 carry_matrix DT。", "低。注意保持 hold/tenor factor 顺序。", ""],
        ["8", "Curve Trade 工作区", "目标右下同屏：tabs、structure segmented buttons、held DV01、legs 表、组合图、summary。", "当前 Curve Trade 是第二个完整 section，通常在首屏下方；首屏不能同时看到 single + curve trade。", "app.R 中 section_card('Curve Trade') 包含 KPI、explanation、DT、两个 plots。", "布局差距大，当前更像纵向报告。", "重组为右下单卡；保留 DT 明细，新增 summary compact card，组合图缩小。", "中到高。DT 宽度、列名、P&L 颜色和响应式 overflow 需单独处理。", ""],
        ["9", "状态与计算", "目标左侧 status 直接显示 Calculation complete、进度条 100%、Run ID。", "当前 carry/trade progress 分散在不同 control_section 内；截图中 Carry complete 可见但不统一。", "status reactiveValues，progress_box('carry') 和 progress_box('trade')。", "状态逻辑已有，表达和位置不一致。", "新增 Carry 页状态卡，整合 carry/trade 状态；Run ID 可用本地时间生成或显示 NA。", "低。不要让输入变更时误清旧结果。", ""],
        ["10", "不可用字段", "目标含 Save Workspace、Market Data Live、Bloomberg BVAL、Curve Quality Good 等生产字段。", "当前显示 Local RDS analytics，本地数据源真实；无 Save Workspace。", "README/sampleUI 说明不可支持模块应保留并标注 unavailable。", "目标中部分字段当前数据不支持。", "保留位置时标 Local RDS / Unavailable / NA；不加入假实时戳。", "低。数据真实性优先于视觉填空。", ""],
        ["11", "响应式/溢出", "目标一屏密度高，1512px 宽下 sidebar + 主区无重叠。", "当前截图右侧图表可见，但页面滚动长；sidebar 内部滚动，selectize 已挂 body。", "styles.css 已设置 sidebar overflow、selectize z-index。", "需要从纵向滚动报告改为首屏工作台。", "新增 Carry 专属 grid 和媒体查询；不改全局 dashboard grid。", "中。避免误伤 Forward/Curve 页面。", ""],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = [0.35, 0.75, 1.35, 1.45, 1.35, 1.0, 1.35, 1.0, 0.9]
    set_table_width(table, widths)
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(header)
        style_run(r, bold=True, size=7.1)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            style_run(r, size=6.8)
            if i == 5 and "大" in value:
                r.font.color.rgb = RGBColor.from_string(RED)
    doc.add_paragraph()


def add_plan_table(doc):
    headers = ["阶段", "修改重点", "验收方式"]
    rows = [
        ["1. 证据基线", "保留目标图、当前运行截图、截图条件；把当前截图状态作为改造前基线。", "Word 中两张截图和采集条件完整。"],
        ["2. Carry shell", "新增 Carry 页面专属 class、标题/meta、6 KPI strip、左侧统一 sidebar。", "运行 app 后首屏结构接近目标图，其他 tabs 不变。"],
        ["3. 图表重排", "把 3M carry/roll/total by tenor 做成主 decomposition 图；Spot Curve 和 Heatmap 调整位置与尺寸。", "Plotly legend/modebar 不压标题，tenor 顺序正确。"],
        ["4. Curve Trade 卡", "把 trade inputs、legs DT、portfolio chart、summary 整合到右下工作区。", "点击 Calculate Curve Trade 后表格和 summary 更新。"],
        ["5. 测试与截图", "运行 tests/run_tests.R、HTTP smoke、浏览器截图对比。", "测试通过，截图与目标布局差距只剩可接受 P3 polish。"],
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [1.0, 5.2, 2.8])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        r = cell.paragraphs[0].add_run(header)
        style_run(r, bold=True, size=8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            r = cell.paragraphs[0].add_run(value)
            style_run(r, size=8.2)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Carry & Roll 参考 UI 对比与实现建议")
    style_run(run, bold=True, color=INK, size=20)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("基于 reference-ui-implementation 工作流；目标图为 sampleUI/04_carry_roll.png，当前截图来自真实运行 Shiny 页面。")
    style_run(run, color="66788C", size=9)

    add_callout(doc, "结论摘要", [
        "当前 Carry & Roll 的计算逻辑和 applied-result 模式可保留；主要工作是把纵向报告式布局改成目标图的一屏交易台布局。",
        "目标图中的 Live Bloomberg、Save Workspace、实时 timestamp 等生产字段当前本地 RDS 不支持，文档建议保留位置但显示 Local RDS / Unavailable / NA。",
        "执行时应使用 Carry 页专属 CSS 和 helper，避免改坏 Curve Explorer、History、Forward、Diagnostics。"
    ], fill="EEF6FF")

    add_heading(doc, "1. 视觉证据", 1)
    add_body(doc, "目标 UI 截图：用户提供 PNG，与项目内 sampleUI/04_carry_roll.png 一致。", bold=True)
    doc.add_picture(str(TARGET_IMAGE), width=Inches(9.6))
    add_caption(doc, "图 1：目标 Carry & Roll 页面，一屏展示 sidebar、KPI、双图、热力图与 Curve Trade 工作区。")

    add_body(doc, "当前运行 UI 截图：真实启动 Shiny 后，使用 Chrome DevTools 协议进入 Carry & Roll 并点击计算按钮采集。", bold=True)
    doc.add_picture(str(CURRENT_IMAGE), width=Inches(9.6))
    add_caption(doc, "图 2：当前 Carry & Roll 页面 applied-state 截图，显示旧版模块 tabs、Single Trade section 和较长 sidebar。")

    add_key_value_table(doc, [
        ("URL", "http://127.0.0.1:7855"),
        ("Page / Tab", "Carry & Roll"),
        ("关键输入", "Zero-rate snapshot；Curve = USD UNITED STATES OIS；Fit method = Nelson-Siegel；Trade start = 0Y；Trade end = 5Y；Hold = 3M；Direction = Receive Fixed；DV01 per bp = 10000；Trade structure = Steepener 默认值。"),
        ("是否点击 Apply/Calculate", "已点击 Carry & Roll tab；已点击 Calculate Carry & Roll；已点击 Calculate Curve Trade。截图可见 Carry 状态 Complete / 100%。"),
        ("截图时间", "2026-06-20 01:30 左右，Asia/Shanghai，本地 Shiny 运行环境。"),
        ("截图状态", "当前截图来自真实运行页面；不是代码结构图。可见首屏顶部，Curve Trade 结果区域在页面下方/折叠滚动区域。"),
    ])

    add_heading(doc, "2. 目标 UI 模块地图", 1)
    for item in [
        "Navbar：深 navy 顶栏，Carry & Roll active，右侧状态/通知/用户入口。",
        "Sidebar：单一控制栏，从数据源、曲线、日期、期限范围到交易方向、DV01、计算按钮和状态。",
        "Header：Carry & Roll Analysis 标题加 Source / Curve / Effective Date / Hold Basis metadata。",
        "KPI：Carry、Roll、Total、P&L、Trade Mode、Curve Quality 六张横向卡片。",
        "Charts：上方左侧为多期限 carry/roll/total decomposition，右侧为 spot curve。",
        "Heatmap：左下矩阵以 Hold Period 为行、Tenor 为列，用红-白-绿表达 total bp。",
        "Curve Trade：右下同屏工作区，包含结构选择、held DV01、legs 表、组合图和 summary。",
        "Status：sidebar 底部显示 calculation complete、进度条、Run ID。"
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. 当前 UI 与代码盘点", 1)
    add_key_value_table(doc, [
        ("入口", r"C:\Users\PC\Desktop\R_git\YieldCurve_Shiny\run_app.R；真实 app 实现在 app.R。"),
        ("当前 UI 结构", "Carry tab 使用 dashboard_page；sidebar 内分 Single trade source、Single trade、Curve trade structure；主区先 module_tabs，再 Single Trade section 和 Curve Trade section。"),
        ("核心 outputs", "carry_banner、carry_value、roll_value、total_value、carry_direction_value、carry_component_plot、carry_spot_plot、carry_stacked_plot、carry_heatmap、carry_matrix、trade_leg_table、trade_leg_pnl_plot、trade_component_plot。"),
        ("reactive 方式", "applied_carry() 和 applied_trade() 保存点击 Calculate 后的结果；输入变更只标记 pending，不立即覆盖旧结果。这个模式应保留。"),
        ("计算引擎", "calculate_carry_roll()、build_carry_matrix()、curve_trade_legs()、calculate_curve_trade() 已支持所需 carry/roll/DV01/P&L 口径。"),
        ("样式入口", "www/styles.css 已有 dashboard、metric、plot_card、sidebar、selectize z-index 等通用样式；建议新增 Carry 页专属 class 而非全局覆盖。"),
        ("已验证", "tests/run_tests.R 已通过；当前截图也证明 Shiny 可以在 http://127.0.0.1:7855 真实运行。"),
    ])

    add_heading(doc, "4. 逐项差距表", 1)
    add_gap_table(doc)

    add_heading(doc, "5. 推荐落地计划", 1)
    add_plan_table(doc)

    add_heading(doc, "6. 测试与验收清单", 1)
    for item in [
        r"运行 C:\Program Files\R\R-4.5.2\bin\Rscript.exe C:\Users\PC\Desktop\R_git\YieldCurve_Shiny\tests\run_tests.R。",
        "启动 tests/smoke_app.R，确认 HTTP 200，并用浏览器进入 Carry & Roll。",
        "点击 Calculate Carry & Roll 和 Calculate Curve Trade 后截图，与目标 PNG 对比。",
        "重点检查：sidebar dropdown 不被裁剪、Plotly modebar 不压标题、KPI 文案不溢出、DT 表格横向滚动可控、1512px 宽度首屏无重叠。",
        "目标图无法真实支持的字段不得伪造；必须显示 Local RDS、Unavailable 或 NA。",
        "遵守删除限制：禁止批量删除文件或目录；若必须删除，只能一次删除一个明确路径。"
    ]:
        add_bullet(doc, item)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
