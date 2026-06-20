# -*- coding: utf-8 -*-
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
TARGET_IMG = BASE / "sampleUI" / "05_diagnostics.png"
CURRENT_IMG = BASE / "diagnostics_current_ui.png"
OUT = BASE / "Diagnostics_UI_Comparison_Plan.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9E2EC"):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn("w:" + edge))
        if element is None:
            element = OxmlElement("w:" + edge)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_east_asia_font(style, font_name):
    rpr = style._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style._element.append(rpr)
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def add_heading(doc, text, level=1):
    doc.add_paragraph(text, style=f"Heading {level}")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(85, 101, 117)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    set_east_asia_font(styles["Normal"], "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        set_east_asia_font(styles[name], "Microsoft YaHei")
    styles["Title"].font.size = Pt(20)
    styles["Title"].font.color.rgb = RGBColor(11, 37, 69)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.size = Pt(11)
    styles["Heading 2"].font.color.rgb = RGBColor(31, 77, 120)
    styles["Heading 3"].font.size = Pt(10)
    styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        paragraph = styles[style_name].paragraph_format
        paragraph.space_after = Pt(4)
        paragraph.line_spacing = 1.08
    return doc


def build():
    doc = setup_document()
    screenshot_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    title = doc.add_paragraph(style="Title")
    title.add_run("Diagnostics Tab 目标 UI 对比与下一步建议")
    subtitle = doc.add_paragraph()
    subtitle.add_run("项目：YieldCurve_Shiny | 页面：Diagnostics | 输出：参考 UI 对比稿与落地建议").bold = True
    subtitle.paragraph_format.space_after = Pt(8)

    add_heading(doc, "1. 视觉证据", 1)
    add_heading(doc, "1.1 目标 UI 截图", 2)
    doc.add_picture(str(TARGET_IMG), width=Inches(9.6))
    add_caption(doc, "目标 UI：YieldCurve_Shiny/sampleUI/05_diagnostics.png")
    add_heading(doc, "1.2 当前运行 UI 截图", 2)
    doc.add_picture(str(CURRENT_IMG), width=Inches(9.6))
    add_caption(doc, "当前运行 UI：Apply Curve 后进入 Diagnostics tab 的真实页面截图。")

    add_heading(doc, "1.3 当前截图采集条件", 2)
    conditions = [
        ("URL", "http://127.0.0.1:10021/"),
        ("Page / Tab", "Diagnostics"),
        ("关键输入", "Curve Explorer 默认数据源与默认曲线；点击 Apply Curve 后切换至 Diagnostics"),
        ("是否点击 Apply/Calculate", "已点击 Apply Curve；Diagnostics 本身不触发重新计算"),
        ("截图时间", screenshot_time),
        ("截图状态", "成功；Chrome headless + DevTools Protocol 截图"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    for idx, text in enumerate(["字段", "值"]):
        shade_cell(table.rows[0].cells[idx], "F2F4F7")
        set_cell_text(table.rows[0].cells[idx], text, bold=True, size=8)
    for key, value in conditions:
        row = table.add_row().cells
        set_cell_text(row[0], key, bold=True, size=8)
        set_cell_text(row[1], value, size=8)

    add_heading(doc, "2. 目标 UI 模块地图", 1)
    add_bullets(doc, [
        "Navbar：深色顶部导航，Diagnostics active，右侧 Market Data live、刷新时间、通知和用户头像。",
        "Sidebar：Diagnostics Policy rail，包含 Model Policy、Data Source Policy、Proxy Analysis、Units Policy、Validation Checks、Refresh Local RDS、Local RDS Status。",
        "Header：Diagnostics 标题与 “Last successful Curve Explorer apply result” 副标题。",
        "KPI：5 张质量卡：Data Freshness、Missing Points、Fit RMSE (Spline)、Proxy Flag、Unit Check；带 OK/Warning 图标和阈值说明。",
        "Main：Diagnostics Summary 表 + Fit Residual Diagnostics 图。",
        "Bottom：Input Points & Missing Tenors 矩阵 + About Diagnostics 说明卡 + footer 时间/auto-refresh 状态。",
    ])

    add_heading(doc, "3. 当前 UI 盘点", 1)
    add_bullets(doc, [
        "当前 Diagnostics 仍使用通用 dashboard_page + side_panel + module_tabs 布局，视觉层级和目标图差距较大。",
        "Sidebar 当前只有 Model policy 三个 rail-note，缺少目标图中的 Policy 控件、validation checks、refresh/status 等结构。",
        "Main 当前有 4 个 metric_card：Data Source、Proxy Flag、Fit RMSE、Tenor Points；缺少 Data Freshness、Missing Points、Unit Check 等质量语义。",
        "当前主区没有 residual chart；只有 Fit Diagnostics DT 表和 Input Points DT 表。",
        "当前输出依赖 last applied Curve Explorer result，符合目标图“Diagnostics do not trigger recalculation”的数据原则。",
        "当前代码入口集中在 app.R 的 Diagnostics tab 和 output$diag_* / output$diagnostics_table / output$input_points。",
    ])

    add_heading(doc, "4. 逐项差距表", 1)
    columns = ["编号", "模块", "目标 UI 表达", "当前 UI 截图状态", "当前代码状态", "差距判断", "推荐修改", "实现风险", "用户补充/决策"]
    rows = [
        ["1", "Navbar / Header", "目标为深色 nav + Diagnostics active + Market Data live + refreshed time。标题区简洁。", "当前 nav 已深色，active 状态可见；右侧为 Local RDS loaded 文案；标题区仍是通用页面 header。", "navbarPage + dashboard_page header。", "中等差距。", "保持真实 Local RDS 口径，不伪造 live；标题副标题改为目标文案。", "低。", "确认是否保留 Local RDS 而非 Market Data live。"],
        ["2", "Sidebar Policy", "左侧是完整 Diagnostics Policy 控制台，多组输入和状态。", "当前只有三个说明 note，缺少 policy controls。", "side_panel 中 control_section(\"Model policy\")。", "大差距。", "重建为目标 rail：Model Policy、Data Source Policy、Proxy Analysis、Units Policy、Validation Checks、Refresh、Local RDS Status。", "中：部分控件先作为 policy UI，不一定立即改变计算。", "确认控件是否先做 UI-only。"],
        ["3", "KPI Cards", "5 张质量 KPI，包含 freshness、missing、RMSE、proxy、unit check。", "当前 4 张基础指标，偏技术字段。", "diag_source、diag_proxy、diag_rmse、diag_points。", "大差距。", "改为 5 KPI；可从 applied_curve 派生 missing/tenor count/RMSE/unit policy，freshness 用本地 snapshot 时间或 NA。", "中：freshness/live source 不能伪造。", "确认 freshness 显示 Local RDS timestamp/NA。"],
        ["4", "Diagnostics Summary", "目标为字段、值、单位、来源、状态五列表，带 OK/Warn/N/A 图标。", "当前没有 summary 表，只显示 diagnostics raw table。", "无 summary 输出。", "大差距。", "新增 renderDT 或 HTML table 汇总 Curve、Dates、Source、Proxy、Tenor Count、RMSE 等。", "中：状态规则需定义阈值。", "确认 warning/fail 阈值。"],
        ["5", "Residual Chart", "目标有 Fit Residual Diagnostics bar chart，显示 warn/fail 阈值线。", "当前没有 chart。", "curve_fit()$diagnostics 含 residual_bp，可直接绘图。", "大差距但数据可支持。", "新增 Plotly residual bar：x tenor_label、y residual_bp、阈值线 +/-warn、+/-fail。", "低到中：legend/modebar/margins 需验收。", "默认 warn=1bp fail=2.5bp？"],
        ["6", "Input Points Matrix", "目标是 tenor-by-row 矩阵：Expected、Available、Yield、Source、Age。", "当前是 Input Points DT 明细表，分页显示 raw numeric tenor/rate。", "output$input_points DT。", "大差距。", "将明细改为矩阵视觉；tenor 使用 1M/3M/…/30Y 标签；missing 显示红色 X。", "中：source/age 当前未必真实支持。", "Source/Age 不支持时显示 B/NA 或 Local RDS。"],
        ["7", "About Diagnostics", "目标右下说明卡解释不触发重算、检查项、跳转 Policy。", "当前只有左侧 explanation card。", "diagnostics_explanation 简短文本。", "中到大差距。", "新增 About card，保留“不触发重算”原则，说明检查项；跳转可做 disabled/unavailable。", "低。", "确认是否需要 Policy Panel 链接。"],
        ["8", "Unavailable Metadata", "目标不突出 unavailable，占位内容被 policy 和说明吸收。", "当前底部两个大 placeholder 占明显空间。", "unavailable_card 两个。", "中等差距。", "保留数据真实性，但降级到 summary/About 中的 NA/Unavailable 行，不再占主视觉。", "低。", "确认是否保留现金流引擎占位。"],
        ["9", "视觉密度", "目标一屏信息密集，卡片圆角小、边框浅、表格紧凑。", "当前留白大，DT 表格占高度，信息关系弱。", "通用 CSS，Diagnostics 无专属 class。", "大差距。", "添加 diagnostics-page 专属 class 与 CSS，避免误伤其他 tabs。", "中：需多视口截图验收。", "默认按目标 1490px 宽度优化。"],
        ["10", "数据真实性", "目标含 Bloomberg/Live/Refinitiv 等生产字段。", "当前真实口径是 Local RDS / fitted proxy。", "data_loader + applied_curve。", "需要谨慎。", "不伪造 live Bloomberg；显示 Local RDS、NA、Unavailable；支持的字段从 applied_curve 派生。", "低：原则清楚。", "确认文案是否接受 Local RDS。"],
    ]

    diff_table = doc.add_table(rows=1, cols=len(columns))
    diff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(diff_table)
    for idx, col in enumerate(columns):
        shade_cell(diff_table.rows[0].cells[idx], "E8EEF5")
        set_cell_text(diff_table.rows[0].cells[idx], col, bold=True, size=7)
    for row_data in rows:
        cells = diff_table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value, size=6.6)

    add_heading(doc, "5. 下一步推荐", 1)
    add_bullets(doc, [
        "第一步先改 Diagnostics 信息架构，不碰计算核心：新增 diagnostics-page 专属布局、sidebar policy rail、顶部 5 KPI 和 About card。",
        "第二步新增 summary 派生表：从 applied_curve、curve_fit()$diagnostics 和 input/source metadata 生成字段/值/单位/来源/状态。",
        "第三步新增 residual Plotly 图：使用现有 residual_bp，加入 warn/fail 阈值线；图例放上方或下方，避免 modebar 重叠。",
        "第四步把 Input Points 改为 missing-tenor matrix：目标 tenors 固定为 1M、3M、6M、1Y、2Y、3Y、5Y、7Y、10Y、15Y、20Y、30Y。",
        "第五步处理数据真实性：Bloomberg live、Refinitiv、age seconds、policy panel link 等当前无真实数据时显示 Local RDS / NA / Unavailable。",
        "第六步验收：run_tests.R、HTTP smoke、Diagnostics 截图；如果 browser 自动化失败，文档或验收记录中明确写明。",
    ])

    add_heading(doc, "6. 建议的实现顺序", 1)
    add_bullets(doc, [
        "新增 diagnostics helpers：状态图标、KPI 卡、summary row、tenor matrix formatter。",
        "重排 Diagnostics tab UI：sidebar + KPI strip + summary/chart + matrix/about。",
        "新增 residual chart output；保留旧 diagnostics_table/input_points 到隐藏 legacy 或逐步替换。",
        "新增 diagnostics 专属 CSS，限定在 .diagnostics-page，避免影响 Carry & Roll、Curve Explorer。",
        "更新 tests/run_tests.R：覆盖 summary、residual plot、matrix 输出和 applied-result 不重算原则。",
    ])

    add_heading(doc, "7. 明确假设", 1)
    add_bullets(doc, [
        "目标图是 Diagnostics 的唯一视觉准绳，但数据真实性高于视觉填充。",
        "Diagnostics 不触发重新计算，只读取最近一次 Curve Explorer Apply 结果。",
        "当前不能真实支持 Live Bloomberg、Refinitiv、精确 quote age、Policy Panel 跳转时，显示 Local RDS、NA 或 Unavailable。",
        "先交付 Word 对比稿和下一步建议，不直接改 Diagnostics 代码。",
    ])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
